import os
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
import tempfile
import json
from pydub import AudioSegment
import gradio as gr
from faster_whisper import WhisperModel
import zipfile
from docx import Document
import re
from openpyxl.styles import Alignment, Font, PatternFill
import pandas as pd
from tqdm import tqdm
from datetime import datetime
import traceback
import ctranslate2
import csv
import gc
import numpy as np
import ffmpeg
import math
from multiprocessing import Process, Queue

def get_audio_duration(filepath, output_dir="temp_wav_files"):
    """
    Converts an input audio/video file to WAV format with 16kHz sampling rate and mono channel.
    Returns the duration (in seconds) of the audio and the path to the converted WAV file.

    Args:
        filepath (str): Path to the input file (mp3, mp4, webm, mkv, etc.).
        output_dir (str): Directory to save the converted WAV file.

    Returns:
        tuple: (duration_in_seconds (float), wav_filepath (str))
    """
    try:
        # Ensure the output directory exists
        if not filepath.endswith(".wav"):
            
   

            # Generate a temporary output path for the WAV file
            filename = os.path.splitext(os.path.basename(filepath))[0]
            wav_filepath = os.path.join(output_dir, f"{filename}.wav")

            # Convert the input file to WAV format with 16kHz and mono
            ffmpeg.input(filepath).output(
                wav_filepath,
                format="wav",
                ac=1,  # mono audio
                ar="16000",  # 16kHz sampling rate
            ).run(overwrite_output=True)
        else:
            wav_filepath=filepath

        # Get the duration of the converted file using ffprobe
        probe = ffmpeg.probe(wav_filepath)
        duration = float(next(
            (stream["duration"] for stream in probe["streams"] if "duration" in stream),
            0
        ))

        # Return the duration and path to the WAV file
        return duration, wav_filepath

    except Exception as e:
        print(f"Error in convert_to_wav_and_get_duration: {e}")
        traceback.print_exc()
        return None, None

def format_timestamp(seconds):
    hrs, secs = divmod(seconds, 3600)
    mins, secs = divmod(secs, 60)
    millis = int((secs % 1) * 1000)
    return f"{int(hrs):02}:{int(mins):02}:{int(secs):02},{millis:03}"

#dataframe追加
def parse_srt_c(srt_content):
    pattern = r'(\d+)\n(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})\n(.*?)\n\n'
    matches = re.findall(pattern, srt_content, re.DOTALL)
    
    subtitles = []
    for match in matches:
        subtitles.append({
            'ID': int(match[0]),
            'Start': match[1],
            'End': match[2],
            'Text': match[3].replace('\n', ' ')
        })
    
    return subtitles

def dataframe_to_html_table(df):
    return df.to_html(index=False)

# SRTファイルからExcelファイルを作成する関数
def create_excel_from_srt_c(srt_content, input_file_name):
    excel_file_name = f"{input_file_name}_srt.xlsx"
    english_subtitles = parse_srt_c(srt_content)

    data = []
    for eng in english_subtitles:
        data.append({
            'ID': eng['ID'],
            'Start': eng['Start'],
            'End': eng['End'],
            'English Subtitle': eng['Text']
        })

    df = pd.DataFrame(data)
    timestamp_patch = datetime.now().strftime("%Y%m%d%H%M%S")
    temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp_patch}")
    os.makedirs(temp_dir, exist_ok=True)
    
    excel_file_path = os.path.join(temp_dir, excel_file_name)
    print(f"Temporary directory: {temp_dir}")
    
    with pd.ExcelWriter(excel_file_path, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Subtitles')
        workbook = writer.book
        worksheet = writer.sheets['Subtitles']

        column_widths = {'A': 7, 'B': 25, 'C': 25, 'D': 90, 'E': 90}
        for column, width in column_widths.items():
            worksheet.column_dimensions[column].width = width

        for row in worksheet.iter_rows(min_row=2, max_row=len(df) + 1):
            for cell in row:
                if cell.column_letter == 'A':
                    cell.alignment = Alignment(horizontal='right', vertical='center')
                elif cell.column_letter in ['B', 'C']:
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                elif cell.column_letter in ['D', 'E']:
                    cell.alignment = Alignment(horizontal='left', vertical='center')

        for row in worksheet.iter_rows(min_row=2, max_row=len(df) + 1):
            worksheet.row_dimensions[row[0].row].height = 30

        header_font = Font(bold=True)
        for cell in worksheet["1:1"]:
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center')
            cell.fill = PatternFill(start_color="DAEEF3", end_color="DAEEF3", fill_type="solid")

    
    return excel_file_path, df

'''def exe_for_gradio(srt_content, input_file_name='Noname'):
    excel_filepath, df_display = create_excel_from_srt_c(srt_content, input_file_name)
    return df_display'''


def transcribe(queue,File, Model, Computing, Lang, BeamSize, VadFilter, device):

    '''global model
    del model
    gc.collect()'''  
    
    FileName = File
    if Lang == "日本語":
        Lang = "ja"
    else:
        Lang = "en"

    with open("replacements_tr.csv", newline="", encoding="utf-8") as f:
        reader = csv.reader(f)
        replacements = [row for row in reader]

    with open("dot_replacements_tr.csv", newline="", encoding="utf-8") as f:
        reader = csv.reader(f)
        dot_replacements = [row for row in reader]
        #print(dot_replacements)


    if device=="cuda":
        #print(ctranslate2.get_cuda_device_count())
        if ctranslate2.get_cuda_device_count() == 0:
             queue.put(("error","Cudaが選択されていますが、利用できません。"))
             
    
    total_duration,File = get_audio_duration(File)
    model = WhisperModel(Model, device=device, compute_type=Computing)
    print(f"using:{device}")



    segments, _ = model.transcribe(File, word_timestamps=True, beam_size=BeamSize, initial_prompt="Hello, I am Scott.", language=Lang, vad_filter=VadFilter)

        #error_message = f"文字起こし中にエラーが発生しました: {e}"
        #return error_message, "", "", [], [], "", "", "", "", "",""

    
    if isinstance(total_duration, str):  # get_audio_duration関数がエラーメッセージを返した場合
        queue.put(("error",None))
    
    words_data = []

    try:
        # Initialize tqdm progress bar
        progress_bar = tqdm(total=total_duration, unit="s", position=0, leave=True, desc="処理進行状況")
        last_update_time = 0

        # 初期更新
        queue.put(("progress",0))
        progress_bar.update(0)

        for segment in segments:
            for word in segment.words:
                word_info = {
                    "start": word.start,
                    "end": word.end,
                    "word": word.word
                }
                words_data.append(word_info)
            
            # Update progress at a reasonable frequency
            segment_progress = segment.end - last_update_time
            if segment_progress >= 1.0:  # Update every 1 second
                progress_bar.update(segment_progress)
                queue.put(("progress", (segment.end / total_duration)))
                last_update_time = segment.end        
                progress_bar.set_postfix({"progress": segment.end / total_duration * 100})

        progress_bar.update(total_duration - last_update_time)
        queue.put(("progress", 1))  # Ensure progress is 100% at the end
        # Close tqdm progress bar
        progress_bar.close()
        
    except Exception as e:
        print(f"進捗バー更新中にエラーが発生しました(ループ処理中のエラー): {e}")
        traceback.print_exc()
        queue.put(("error",None))
    finally:
        if os.path.exists(File):
            os.remove(File)
           
    #torch.cuda.empty_cache() 

    try:
        def merge_words(words_data):
            """
            辞書データの単語を結合し、開始時刻と終了時刻を適切に調整する。
            
            Args:
                words_data (list): [{'start': float, 'end': float, 'word': str}, ...]の形式のデータ。
            
            Returns:
                list: 修正済みの単語リスト [{'start': float, 'end': float, 'word': str}, ...]
            """
            merged_data = []
            append_punctuations = [".", ",", "?", "!", ";", ":"]  # 結合対象の記号

            def normalize_floats(wordstamps):
                """
                リスト内の辞書データを正規化（np.float64 → float）。
                """
                normalized = []
                for wordstamp in wordstamps:
                    normalized.append({
                        'start': float(wordstamp['start']),
                        'end': float(wordstamp['end']),
                        'word': wordstamp['word']
                    })
                return normalized
            words_data = normalize_floats(words_data)        

            for word_data in words_data:
                word = word_data["word"]  # 空白を保持
                if word.startswith(tuple(append_punctuations)) and merged_data:
                    # 直前の単語に結合 (スペースを保持)
                    merged_data[-1]["word"] += word
                    # 終了時刻を更新
                    merged_data[-1]["end"] = word_data["end"]
                else:
                    # 新しい単語として追加
                    merged_data.append(word_data)
            
            return merged_data 
               
        words_data=merge_words(words_data)
        
        #print("ファイル処理-Dot変換")
        for word_info in words_data:
            for original, replacement, _, _, _ in replacements:
                original=re.escape(original)
                original=rf"\b{original}\b"
                word_info['word']= re.sub(original,replacement,word_info['word'])

        for word_info in words_data:
            for dot_original, dot_replacement, _, _, _ in dot_replacements:
                escaped_dot_original = re.escape(dot_original) 
                final_dot_original = rf"\b{escaped_dot_original}"
                word_info["word"] = re.sub(final_dot_original,dot_replacement,word_info["word"])
               

       


        # 前処理: words_data内の各wordの中から★を削除する
        cleaned_words_data = []
        for word_info in words_data:
            cleaned_word_info = {
                "start": word_info["start"],
                "end": word_info["end"],
                "word": word_info["word"].replace("[dot]", ".")
            }
            cleaned_words_data.append(cleaned_word_info)

        input_file_name = os.path.splitext(os.path.basename(File))[0]
        
        timestamp_patch = datetime.now().strftime("%Y%m%d%H%M%S")
        temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp_patch}")
        os.makedirs(temp_dir, exist_ok=True)
        
        json_output_file_name = f"{input_file_name}.json"
        json_output_path = os.path.join(temp_dir, json_output_file_name)
        # JSONファイルへの書き込み
        with open(json_output_path, 'w', encoding='utf-8') as f:
           json.dump(cleaned_words_data, f, ensure_ascii=False, indent=4)
           
        # 書き込んだJSONデータの表示（デバッグ用）
        json_content = json.dumps(cleaned_words_data, ensure_ascii=False, indent=4)

        srt_entries = []
        entry_number = 1
        segment_text = ""
        segment_start = None
        segment_end = None

        for word_info in words_data:
            if segment_start is None:
                segment_start = word_info["start"]
            
            segment_text += word_info["word"]
            segment_end = word_info["end"]
            
            if word_info["word"].endswith('.') or word_info["word"].endswith('?'):
                srt_entries.append({
                    "number": entry_number,
                    "start": segment_start,
                    "end": segment_end,
                    "text": segment_text.strip()
                })
                entry_number += 1
                segment_text = ""
                segment_start = None

        if segment_text.strip():
            srt_entries.append({
                "number": entry_number,
                "start": segment_start,
                "end": segment_end,
                "text": segment_text.strip()
            })

        srt_output_file_name = f"{input_file_name}.srt"
        srt_output_path = os.path.join(temp_dir, srt_output_file_name)

        with open(srt_output_path, 'w', encoding='utf-8') as f:
            for entry in srt_entries:
                start_time = format_timestamp(entry["start"])
                end_time = format_timestamp(entry["end"])
                text = entry['text'].replace("[dot]", ".")
                f.write(f"{entry['number']}\n{start_time} --> {end_time}\n{text}\n\n")

        with open(srt_output_path, 'r', encoding='utf-8') as f:
            srt_content = f.read()

        txt_nr_content = ""
        for word_info in words_data:
            if not txt_nr_content:
                txt_nr_content += word_info['word'].lstrip()
            else:
                txt_nr_content += word_info['word']

        txt_nr_output_file_name = f"{input_file_name}_NR.txt"
        txt_nr_output_path = os.path.join(temp_dir, txt_nr_output_file_name)
        with open(txt_nr_output_path, 'w', encoding='utf-8') as f:
            txt_nr_content = txt_nr_content.replace("[dot]","")
            f.write(txt_nr_content)

        txt_r_content = ""
        previous_word_end = 0
        is_first_word = True
        for word in words_data:
            if is_first_word or txt_r_content.endswith("\n"):
                txt_r_content += word['word'].strip()
            else:
                txt_r_content += word['word']
        
            if "." in word['word']:
                if word['start'] - previous_word_end >= 0.5:
                    txt_r_content += "\n"
                previous_word_end = word['end']
            is_first_word = False

        txt_r_output_file_name = f"{input_file_name}_R.txt"
        txt_r_output_path = os.path.join(temp_dir, txt_r_output_file_name)

        with open(txt_r_output_path, 'w', encoding='utf-8') as f:
            txt_r_content = txt_r_content.replace("[dot]", ".")
            f.write(txt_r_content)

        # srtファイルからワードファイルへ変換
        doc_srt = Document()

        srtdoc_output_file_name = f"{input_file_name}_srt.docx"
        srtdoc_output_path = os.path.join(temp_dir, srtdoc_output_file_name)

        with open(srt_output_path, 'r', encoding='utf-8') as file:
            lines = file.readlines()

        subtitle_number = None
        timestamp = None
        subtitle_text = []

        for line in lines:
            line = line.strip()
            if line.isdigit():
                # 以前の字幕エントリを追加
                if subtitle_number is not None and subtitle_text:
                    doc_srt.add_paragraph(f'{subtitle_number}')
                    doc_srt.add_paragraph(f'{timestamp}')
                    doc_srt.add_paragraph(' '.join(subtitle_text))
                    doc_srt.add_paragraph()  # 空行で区切る

                subtitle_number = line
                timestamp = None
                subtitle_text = []
            elif '-->' in line:
                timestamp = line
            elif line:
                subtitle_text.append(line)

        if subtitle_number is not None and subtitle_text:
            doc_srt.add_paragraph(f'{subtitle_number}')
            doc_srt.add_paragraph(f'{timestamp}')
            doc_srt.add_paragraph(' '.join(subtitle_text))

        doc_srt.save(srtdoc_output_path)

        ## txt(nr)をdoc変換
        txtdoc_nr = Document()
        txtdoc_nr_output_file_name = f"{input_file_name}_txtnr.docx"
        txtdoc_nr_output_path = os.path.join(temp_dir, txtdoc_nr_output_file_name)

        with open(txt_nr_output_path, 'r', encoding='utf-8') as file:
            lines = file.readlines()

        for line in lines:
            txtdoc_nr.add_paragraph(line)

        txtdoc_nr.save(txtdoc_nr_output_path)

        ## txt(r)をdoc変換
        txtdoc_r = Document()
        txtdoc_r_output_file_name = f"{input_file_name}_txtr.docx"
        txtdoc_r_output_path = os.path.join(temp_dir, txtdoc_r_output_file_name)

        with open(txt_r_output_path, 'r', encoding='utf-8') as file:
            content = file.read()

        paragraph = txtdoc_r.add_paragraph()
        paragraph.add_run(content)  # 改行はそのまま出力
        txtdoc_r.save(txtdoc_r_output_path)



        #xls,df追加

        

        def parse_timestamp(timestamp):
            """SRTのタイムスタンプを秒に変換"""
            time_format = "%H:%M:%S,%f"
            dt = datetime.strptime(timestamp, time_format)
            return dt.hour * 3600 + dt.minute * 60 + dt.second + dt.microsecond / 1_000_000


        def process_srt(file_path):
            """SRTファイルを読み込んで条件分岐"""
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()
            
            # SRTのセグメントを正規表現で分割
            segments = re.split(r"\n\n", content.strip())
            
            for segment in segments:
                # タイムスタンプ部分を抽出
                match = re.search(r"(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})", segment)
                if match:
                    start_time = parse_timestamp(match.group(1))
                    end_time = parse_timestamp(match.group(2))
                    # 差が20秒以上ならfuncAを実行
                    if end_time - start_time > 60:
                        return "60秒以上のセグメントが見つかりました。ピリオド付加をご検討下さい"
                    else:
                        return "長時間(>60秒）のセグメントはありません。"

        segment_info=process_srt(srt_output_path)               

        paths=[input_file_name,json_output_path,srt_output_path,txt_r_output_path,txt_nr_output_path,
                txtdoc_nr_output_path, txtdoc_r_output_path,segment_info]
        queue.put(("result",paths))
        queue.put(("done",None))
        
        #torch.cuda.empty_cache()  # GPUメモリを解放
        #gc.collect()     

    except Exception as e:
        print(f"ファイル処理中にエラーが発生しました: {e}")
        traceback.print_exc()
        queue.put(("error",None))


def run_with_progress(File, Model, Computing, Lang, BeamSize, VadFilter, device,progress=gr.Progress()):
    timestamp_patch = datetime.now().strftime("%Y%m%d%H%M%S")
    temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp_patch}")
    os.makedirs(temp_dir, exist_ok=True)

    queue = Queue()
    process = Process(target=transcribe, args=(queue,File, Model, Computing, Lang, BeamSize, VadFilter, device))
    process.start()

    paths = None
    flag=False
 
    while True:
        message_type, data = queue.get()
        if message_type == "progress":
            progress(data)  # 進捗更新
        elif message_type == "result":
            paths = data  # パスのリストを受信
        elif message_type == "error":
            flag=True
            break 
        elif message_type == "done":
            break
    process.join()
    if flag==True:
        return "","","", "", [], [] ,"", "", "", "", "", None

    
    
    input_file_name=paths[0]
    json_output_path=paths[1]
    srt_output_path=paths[2]
    txt_r_output_path=paths[3]
    txt_nr_output_path=paths[4]
    txtdoc_nr_output_path=paths[5]
    txtdoc_r_output_path=paths[6]
    segment_info=paths[7]

    with open(srt_output_path,"r",encoding='utf-8') as f:
        srt_content=f.read()
    with open(txt_r_output_path,"r",encoding='utf-8') as f:
        txt_r_content=f.read()
    with open(txt_nr_output_path,"r",encoding='utf-8') as f:
        txt_nr_content=f.read()
    
    excel_filepath, df_display = create_excel_from_srt_c(srt_content=srt_content, input_file_name=input_file_name)

    # zipファイルにまとめる(srt,txtr,txtnr)。
    zip_core_file_name = f"{input_file_name}_core.zip"
    zip_core_file_path = os.path.join(temp_dir, zip_core_file_name)

    with zipfile.ZipFile(zip_core_file_path, 'w') as zip_file:
        zip_file.write(json_output_path, os.path.basename(json_output_path))
        zip_file.write(srt_output_path, os.path.basename(srt_output_path))
        zip_file.write(txt_r_output_path, os.path.basename(txt_r_output_path))
        zip_file.write(txt_nr_output_path, os.path.basename(txt_nr_output_path))
        
    
    # zipファイルにまとめる(doc)。
    zip_doc_file_name = f"{input_file_name}_office_en.zip"
    zip_doc_file_path = os.path.join(temp_dir, zip_doc_file_name)

    with zipfile.ZipFile(zip_doc_file_path, 'w') as zip_file:
        zip_file.write(excel_filepath, os.path.basename(excel_filepath))
        zip_file.write(txtdoc_nr_output_path, os.path.basename(txtdoc_nr_output_path))
        zip_file.write(txtdoc_r_output_path, os.path.basename(txtdoc_r_output_path))

    
    
    main_files = [
        json_output_path,
        srt_output_path,
        txt_nr_output_path,
        txt_r_output_path,
        zip_core_file_path
    ]
    
    zip_doc_file_path = os.path.join(temp_dir, zip_doc_file_name)

    doc_files = [excel_filepath, txtdoc_nr_output_path, txtdoc_r_output_path, zip_doc_file_path]

    ##テーブル##

    df_display=dataframe_to_html_table(df_display)
    df_display=f"""
        <div class="my-table-container">
            {df_display}
        </div>
    """

    html_srt = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 400px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{srt_content}</pre>"""
    html_nr_txt = f"""<pre style="white-space: pre-wrap; overflow-y:auto; height: 400px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{txt_nr_content}</pre>"""
    html_r_txt = f"""<pre style="white-space: pre-wrap; overflow-y:auto; height: 400px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{txt_r_content}</pre>"""

    filename_copy = input_file_name
    srt_dummy_output_path = srt_output_path
    return segment_info,srt_content,txt_nr_content, txt_r_content, main_files, doc_files ,html_srt, html_nr_txt, html_r_txt, filename_copy, srt_dummy_output_path, df_display

'''
テキストエリア①
テキストエリア②
テキストエリア③
ファイル④
ファイル⑤
HTML⑥
HTML⑦
HTML⑧
テキストボックス⑨
テキストボックス⑩
HTML⑪
'''